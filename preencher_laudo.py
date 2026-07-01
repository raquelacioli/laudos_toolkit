#!/usr/bin/env python3
"""
preencher_laudo.py
-------------------
Preenche automaticamente o modelo de laudo pericial (.docx) do Dr. Tiago
a partir de um arquivo .txt com os dados da perícia, e anexa fotos de
documentos como folhas extras ao final do laudo.

Uso:
    python3 preencher_laudo.py MODELO.docx DADOS.txt FOTO1.jpg [FOTO2.jpg ...] -o SAIDA.docx

O TXT pode ser escrito livremente, no estilo:

    NOME DO PERICIADO
    IDADE: 61 ANOS
    ESCOLARIDADE: ENSINO MÉDIO COMPLETO
    ESTADO CIVIL: CASADA
    OCUPAÇÃO HABITUAL: DOMÉSTICA
    TEMPO DE AFASTAMENTO: HÁ VÁRIOS ANOS

    ANAMNESE:
    (texto livre da anamnese)

    EXAME FÍSICO E MENTAL:
    (texto livre / observações adicionais do exame)

    CONCLUSÃO:
    (texto livre da conclusão)

Pequenos erros de digitação nos rótulos (ex: "ANAMENESE", "CONCLUISÃO")
são tolerados via correspondência aproximada.
"""

import sys
import re
import shutil
import zipfile
import unicodedata
import difflib
import argparse
from pathlib import Path
from lxml import etree

sys.path.insert(0, str(Path(__file__).resolve().parent))
try:
    import scan_documento
except ImportError:
    scan_documento = None

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def wtag(tag):
    return f"{{{W}}}{tag}"


# ---------------------------------------------------------------------------
# 1. Parsing do TXT
# ---------------------------------------------------------------------------

CANONICAL_FIELDS = [
    "IDADE",
    "ESCOLARIDADE",
    "ESTADO CIVIL",
    "OCUPACAO HABITUAL",
    "TEMPO DE AFASTAMENTO",
    "ANAMNESE",
    "EXAME FISICO E MENTAL",
    "CONCLUSAO",
    "SEXO",
    "RG",
    "CPF",
]


def strip_accents(s):
    return "".join(
        c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)
    )


def normalize_key(s):
    s = strip_accents(s).upper().strip()
    s = re.sub(r"[^A-Z0-9 ]", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def closest_canonical(key):
    norm = normalize_key(key)
    match = difflib.get_close_matches(norm, CANONICAL_FIELDS, n=1, cutoff=0.6)
    return match[0] if match else norm


LABEL_LINE_RE = re.compile(r"^([A-ZÀ-Úa-zà-ú/ ]{3,40}):\s*(.*)$")


def is_headerlike(line):
    """True se a linha, mesmo sem ':', for muito parecida com um dos campos
    canonicos (ex: 'EXAME FÍSICO E MENTAL' sozinho na linha)."""
    if len(line) > 40:
        return None
    norm = normalize_key(line)
    if not norm:
        return None
    match = difflib.get_close_matches(norm, CANONICAL_FIELDS, n=1, cutoff=0.75)
    return match[0] if match else None


def parse_txt(path):
    """Retorna dict {campo_canonico: valor} e o nome do periciado (se detectado)."""
    raw_lines = Path(path).read_text(encoding="utf-8").splitlines()
    data = {}
    name = None
    current_key = None
    buffer = []

    def flush():
        nonlocal current_key, buffer
        if current_key is not None:
            text = "\n".join(l for l in buffer).strip()
            if text:
                data[current_key] = text
        current_key = None
        buffer = []

    for line in raw_lines:
        stripped = line.strip()
        m = LABEL_LINE_RE.match(stripped) if stripped else None
        header_only = is_headerlike(stripped) if (stripped and not m) else None
        if m:
            flush()
            label, value = m.group(1), m.group(2).strip()
            canon = closest_canonical(label)
            if value:
                data[canon] = value
                current_key = None
            else:
                current_key = canon
                buffer = []
        elif header_only:
            flush()
            current_key = header_only
            buffer = []
        else:
            if current_key is not None:
                if stripped:
                    buffer.append(stripped)
            elif stripped and name is None and not data:
                # primeira linha solta antes de qualquer rotulo = nome
                name = stripped
    flush()
    return data, name


# ---------------------------------------------------------------------------
# 2. Utilidades de manipulação do XML (content controls / texto simples)
# ---------------------------------------------------------------------------

def get_all_text(el):
    return "".join(t.text or "" for t in el.findall(".//w:t", NS))


def set_run_text(run_el, text):
    t = run_el.find("w:t", NS)
    if t is None:
        t = etree.SubElement(run_el, wtag("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text


def find_sdts(root):
    return root.findall(".//w:sdt", NS)


def sdt_items(sdt):
    sdtPr = sdt.find("w:sdtPr", NS)
    if sdtPr is None:
        return []
    dropdown = sdtPr.find("w:dropDownList", NS)
    if dropdown is None:
        return []
    return [
        li.get(wtag("displayText"))
        for li in dropdown.findall("w:listItem", NS)
    ]


def set_sdt_text(sdt, text):
    """Substitui o conteúdo (sdtContent) de um content control por um texto simples,
    preservando a formatação (rPr) do primeiro run existente."""
    content = sdt.find("w:sdtContent", NS)
    if content is None:
        return False
    runs = content.findall(".//w:r", NS)
    if not runs:
        return False
    first = runs[0]
    set_run_text(first, text)
    # remove runs extras
    for r in runs[1:]:
        r.getparent().remove(r)
    return True


def best_matching_item(value, items):
    """Encontra, dentre os itens de um dropdown, o que melhor corresponde ao valor livre."""
    norm_value = normalize_key(value)
    norm_items = {normalize_key(it or ""): it for it in items if it}
    # 1) match exato
    for norm_it, it in norm_items.items():
        if norm_it == norm_value or norm_it == norm_value + ";":
            return it
    # 2) contém
    for norm_it, it in norm_items.items():
        if norm_value and (norm_value in norm_it or norm_it.replace(";", "") == norm_value):
            return it
    # 3) fuzzy
    match = difflib.get_close_matches(norm_value, list(norm_items.keys()), n=1, cutoff=0.55)
    if match:
        return norm_items[match[0]]
    return None


def preceding_label_text(sdt):
    """Retorna o texto do run de rótulo (plain text) imediatamente anterior ao sdt,
    dentro do mesmo parágrafo."""
    parent = sdt.getparent()
    prev = sdt.getprevious()
    texts = []
    while prev is not None:
        if prev.tag == wtag("r"):
            t = prev.find("w:t", NS)
            if t is not None and t.text:
                texts.insert(0, t.text)
        elif prev.tag == wtag("sdt"):
            break
        prev = prev.getprevious()
    return "".join(texts).strip()


# ---------------------------------------------------------------------------
# 3. Preenchimento dos campos
# ---------------------------------------------------------------------------

def fill_dropdown_by_tag(root, tag_substr, chooser, log):
    """Procura o sdt cujo w:tag/w:alias contenha tag_substr (normalizado) e
    aplica 'chooser(items) -> item_escolhido_ou_None'."""
    tag_substr_norm = normalize_key(tag_substr)
    for sdt in find_sdts(root):
        sdtPr = sdt.find("w:sdtPr", NS)
        if sdtPr is None:
            continue
        tagEl = sdtPr.find("w:tag", NS)
        aliasEl = sdtPr.find("w:alias", NS)
        tag_val = tagEl.get(wtag("val")) if tagEl is not None else ""
        alias_val = aliasEl.get(wtag("val")) if aliasEl is not None else ""
        if tag_substr_norm in normalize_key(tag_val) or tag_substr_norm in normalize_key(alias_val):
            items = sdt_items(sdt)
            chosen = chooser(items)
            if chosen:
                set_sdt_text(sdt, chosen)
                log.append(f"OK  [tag:{tag_substr}] -> \"{chosen.strip()}\"")
                return True
    log.append(f"--  [tag:{tag_substr}] campo nao encontrado")
    return False


def fill_dropdown_by_label(root, label_substr, value, log):
    """Procura o sdt cujo rótulo textual precedente contenha label_substr
    e ajusta seu conteúdo para o item mais próximo de value."""
    label_substr_norm = normalize_key(label_substr)
    for sdt in find_sdts(root):
        label = normalize_key(preceding_label_text(sdt))
        if label_substr_norm in label:
            items = sdt_items(sdt)
            if not items:
                continue
            chosen = best_matching_item(value, items)
            if chosen:
                set_sdt_text(sdt, chosen)
                log.append(f"OK  [{label_substr}] -> \"{chosen.strip()}\"")
                return True
            else:
                log.append(f"??  [{label_substr}] sem correspondencia p/ \"{value}\" (mantido valor original)")
                return False
    log.append(f"--  [{label_substr}] campo (dropdown) nao encontrado no modelo")
    return False


def replace_plain_bold_after_label(root, label_text, value, log):
    """Para campos que sao texto simples em negrito (nao sdt), logo apos um rotulo
    (que pode estar num run simples ou dentro de um content control), substitui
    o primeiro run simples em negrito seguinte ao rotulo, dentro do mesmo paragrafo."""
    found = False
    for p in root.findall(".//w:p", NS):
        children = [c for c in p if c.tag in (wtag("r"), wtag("sdt"))]
        label_idx = None
        for i, c in enumerate(children):
            if label_text in get_all_text(c):
                label_idx = i
                break
        if label_idx is None:
            continue
        for c in children[label_idx + 1:]:
            if c.tag == wtag("r"):
                rpr = c.find("w:rPr", NS)
                is_bold = rpr is not None and rpr.find("w:b", NS) is not None
                if is_bold:
                    set_run_text(c, value)
                    found = True
                    break
        if found:
            break
    if found:
        log.append(f"OK  [{label_text}] -> \"{value}\"")
    else:
        log.append(f"--  [{label_text}] campo (texto simples) nao encontrado no modelo")
    return found


def set_paragraph_text(p, text, log, label):
    """Substitui todo o conteudo de texto de um paragrafo por 'text',
    preservando a formatacao do primeiro run (ou criando um run em negrito)."""
    runs = p.findall("w:r", NS)
    if runs:
        set_run_text(runs[0], text)
        for r in runs[1:]:
            r.getparent().remove(r)
    else:
        r = etree.SubElement(p, wtag("r"))
        rpr = etree.SubElement(r, wtag("rPr"))
        etree.SubElement(rpr, wtag("b"))
        set_run_text(r, text)
    log.append(f"OK  [{label}] preenchido")


def paragraphs_by_heading(root, heading_text):
    """Retorna o indice do paragrafo cujo texto == heading_text, dentro da lista
    de paragrafos do body (nivel raiz, ignora paragrafos dentro de tabelas)."""
    body = root.find("w:body", NS)
    paras = body.findall("w:p", NS)
    for i, p in enumerate(paras):
        if get_all_text(p).strip() == heading_text:
            return i, paras
    return None, paras


def fill_document(root, data, log):
    # --- Tabela "DADOS DO(A) PERICIADO(A)" : campos em dropdown ---
    mapping_dropdown = {
        "SEXO": "Sexo:",
        "ESCOLARIDADE": "Escolaridade:",
        "ESTADO CIVIL": "Estado civil:",
        "OCUPACAO HABITUAL": "Ocupacao habitual",
        "TEMPO DE AFASTAMENTO": "Data declarada de afastamento do trabalho:",
    }
    for canon, label in mapping_dropdown.items():
        if canon in data:
            fill_dropdown_by_label(root, label, data[canon], log)

    # --- Campos em texto simples (negrito) ---
    if "IDADE" in data:
        idade_num = re.sub(r"\D", "", data["IDADE"])
        if idade_num:
            replace_plain_bold_after_label(root, "Idade:", idade_num, log)
    if "RG" in data:
        replace_plain_bold_after_label(root, "RG:", data["RG"], log)
    if "CPF" in data:
        replace_plain_bold_after_label(root, "CPF:", data["CPF"], log)

    # --- ANAMNESE: paragrafo vazio logo apos o cabecalho "3. ANAMNESE" ---
    if "ANAMNESE" in data:
        idx, paras = paragraphs_by_heading(root, "3. ANAMNESE")
        if idx is not None:
            # paragrafo idx+1 = frase fixa introdutoria; idx+2 = onde inserir o texto
            target = None
            for j in range(idx + 1, min(idx + 4, len(paras))):
                if get_all_text(paras[j]).strip() == "":
                    target = paras[j]
                    break
            if target is not None:
                set_paragraph_text(target, data["ANAMNESE"], log, "ANAMNESE")
            else:
                log.append("--  [ANAMNESE] paragrafo de destino nao encontrado")
        else:
            log.append("--  [ANAMNESE] cabecalho '3. ANAMNESE' nao encontrado")

    # --- EXAME FISICO E MENTAL: paragrafo vazio apos o paragrafo com os dropdowns ---
    if "EXAME FISICO E MENTAL" in data:
        idx, paras = paragraphs_by_heading(root, "4. EXAME FÍSICO E MENTAL")
        if idx is not None:
            target = None
            for j in range(idx + 1, min(idx + 4, len(paras))):
                if get_all_text(paras[j]).strip() == "":
                    target = paras[j]
                    break
            if target is not None:
                set_paragraph_text(target, data["EXAME FISICO E MENTAL"], log, "EXAME FISICO E MENTAL")
                if "CICATRIZ" in normalize_key(data["EXAME FISICO E MENTAL"]):
                    def choose_cicatriz(items):
                        for it in items:
                            if it and "trauma referido no item 3" in it and "cirurgia" not in it:
                                return it
                        return None
                    fill_dropdown_by_tag(root, "CAIXA - EXAME FISICO", choose_cicatriz, log)
            else:
                log.append("--  [EXAME FISICO E MENTAL] paragrafo de destino nao encontrado")
        else:
            log.append("--  [EXAME FISICO E MENTAL] cabecalho nao encontrado")

    # --- CONCLUSAO: paragrafo apos "conclusao diagnostica:" ---
    if "CONCLUSAO" in data:
        idx, paras = paragraphs_by_heading(
            root, "Com base na aplicação da metodologia acima descrita, chegou-se à seguinte conclusão diagnóstica:"
        )
        if idx is not None and idx + 1 < len(paras):
            set_paragraph_text(paras[idx + 1], data["CONCLUSAO"], log, "CONCLUSAO")
        else:
            log.append("--  [CONCLUSAO] paragrafo de destino nao encontrado")


# ---------------------------------------------------------------------------
# 4. Execucao principal: unzip -> editar XML -> rezip -> anexar fotos
# ---------------------------------------------------------------------------

def rezip(src_dir, out_path):
    out_path = Path(out_path)
    if out_path.exists():
        out_path.unlink()
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(Path(src_dir).rglob("*")):
            if f.is_file():
                zf.write(f, f.relative_to(src_dir))


def embed_image_relationship(unpacked_dir, image_path):
    """Copia a imagem para word/media/, registra a relationship e o content
    type, e retorna o rId a ser usado no w:drawing."""
    import mimetypes

    unpacked_dir = Path(unpacked_dir)
    media_dir = unpacked_dir / "word" / "media"
    media_dir.mkdir(parents=True, exist_ok=True)

    ext = Path(image_path).suffix.lower().lstrip(".") or "jpg"
    if ext == "jpg":
        ext = "jpeg"
    existing = list(media_dir.glob("image*.*"))
    n = len(existing) + 1
    media_name = f"image_face_{n}.{ext}"
    shutil.copy(image_path, media_dir / media_name)

    rels_path = unpacked_dir / "word" / "_rels" / "document.xml.rels"
    rels_tree = etree.parse(str(rels_path))
    rels_root = rels_tree.getroot()
    R_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    existing_ids = [
        int(r.get("Id")[3:]) for r in rels_root
        if r.get("Id", "").startswith("rId") and r.get("Id")[3:].isdigit()
    ]
    new_id = f"rId{max(existing_ids, default=0) + 1}"
    rel = etree.SubElement(rels_root, f"{{{R_NS}}}Relationship")
    rel.set("Id", new_id)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
    rel.set("Target", f"media/{media_name}")
    rels_tree.write(str(rels_path), xml_declaration=True, encoding="UTF-8", standalone=True)

    ct_path = unpacked_dir / "[Content_Types].xml"
    ct_tree = etree.parse(str(ct_path))
    ct_root = ct_tree.getroot()
    CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
    has_default = any(
        d.get("Extension", "").lower() == ext for d in ct_root.findall(f"{{{CT_NS}}}Default")
    )
    if not has_default:
        mime = mimetypes.guess_type(f"x.{ext}")[0] or "image/jpeg"
        default = etree.SubElement(ct_root, f"{{{CT_NS}}}Default")
        default.set("Extension", ext)
        default.set("ContentType", mime)
        ct_tree.write(str(ct_path), xml_declaration=True, encoding="UTF-8", standalone=True)

    return new_id


def crop_to_ratio(image_path, out_path, target_w=3.0, target_h=4.0):
    """Recorta a foto para a proporcao exata 3x4 (sem distorcer), centralizando
    no rosto quando possivel (deteccao facial), ou no centro da imagem caso
    contrario. Tambem aplica um leve realce de nitidez/contraste."""
    import cv2
    import numpy as np

    img = cv2.imread(str(image_path))
    if img is None:
        raise RuntimeError(f"nao foi possivel ler {image_path}")
    h, w = img.shape[:2]
    target_ratio = target_w / target_h  # largura/altura desejada

    cur_ratio = w / h
    if cur_ratio > target_ratio:
        # imagem mais larga que o alvo -> corta as laterais
        new_w = int(h * target_ratio)
        cx = w // 2
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
            faces = cascade.detectMultiScale(gray, 1.1, 5)
            if len(faces):
                fx, fy, fw, fh = max(faces, key=lambda f: f[2] * f[3])
                cx = fx + fw // 2
        except Exception:
            pass
        x0 = max(0, min(w - new_w, cx - new_w // 2))
        crop = img[:, x0:x0 + new_w]
    else:
        # imagem mais alta que o alvo -> corta em cima/embaixo
        new_h = int(w / target_ratio)
        cy = h // 2
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
            faces = cascade.detectMultiScale(gray, 1.1, 5)
            if len(faces):
                fx, fy, fw, fh = max(faces, key=lambda f: f[2] * f[3])
                # centraliza um pouco acima do centro do rosto (deixa espaco p/ ombros)
                cy = fy + int(fh * 0.7)
        except Exception:
            pass
        y0 = max(0, min(h - new_h, cy - new_h // 3))
        crop = img[y0:y0 + new_h, :]

    lab = cv2.cvtColor(crop, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8, 8))
    l = clahe.apply(l)
    crop = cv2.cvtColor(cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR)
    blur = cv2.GaussianBlur(crop, (0, 0), sigmaX=2)
    crop = cv2.addWeighted(crop, 1.3, blur, -0.3, 0)

    cv2.imwrite(str(out_path), crop, [cv2.IMWRITE_JPEG_QUALITY, 95])


def add_face_photo(root, unpacked_dir, image_path, log):
    """Insere a foto de rosto do(a) periciado(a) na celula reservada para a
    foto, ao lado do rotulo 'DADOS DO(A) PERICIADO(A)' (celula seguinte, na
    mesma linha da tabela). O tamanho da foto e 3x4cm. Essa celula deve ter
    sido alargada no modelo (no Word) para caber a foto -- caso contrario,
    cai de volta para o posicionamento flutuante no canto superior direito."""
    rId = embed_image_relationship(unpacked_dir, image_path)

    # localiza a linha da tabela (w:tr) que contem o rotulo
    target_tr = None
    for tr in root.findall(".//w:tr", NS):
        if "DADOS DO(A) PERICIADO" in get_all_text(tr):
            target_tr = tr
            break
    if target_tr is None:
        log.append("--  [foto de rosto] linha 'DADOS DO(A) PERICIADO(A)' nao encontrada")
        return False

    tcs = target_tr.findall("w:tc", NS)
    if len(tcs) < 2:
        log.append("--  [foto de rosto] celula de foto nao encontrada (linha sem 2a coluna)")
        return False
    photo_tc = tcs[1]

    # pega a largura da celula para decidir se ha espaco reservado (> 1.5cm ~ 850 dxa)
    tcPr = photo_tc.find("w:tcPr", NS)
    tcW = tcPr.find("w:tcW", NS) if tcPr is not None else None
    width_dxa = int(tcW.get(wtag("w"))) if tcW is not None and tcW.get(wtag("w")) else 0

    if width_dxa < 850:
        log.append(f"--  [foto de rosto] celula reservada muito estreita ({width_dxa} dxa) -- alargue a coluna no Word para caber a foto")
        return False

    # pega (ou cria) o primeiro paragrafo dentro da celula, mesmo se estiver
    # dentro de um content control (w:sdt/w:sdtContent)
    p = photo_tc.find(".//w:p", NS)
    if p is None:
        p = etree.SubElement(photo_tc, wtag("p"))

    # limpa runs existentes no paragrafo (mantem pPr se houver)
    for r in p.findall("w:r", NS):
        p.remove(r)

    r = etree.SubElement(p, wtag("r"))
    rPr = etree.SubElement(r, wtag("rPr"))
    etree.SubElement(rPr, wtag("noProof"))
    drawing = etree.SubElement(r, wtag("drawing"))

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    inline = etree.SubElement(drawing, f"{{{WP}}}inline")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")

    extent = etree.SubElement(inline, f"{{{WP}}}extent")
    extent.set("cx", "1080000")
    extent.set("cy", "1440000")

    effectExtent = etree.SubElement(inline, f"{{{WP}}}effectExtent")
    for a in ("l", "t", "r", "b"):
        effectExtent.set(a, "0")

    docPr = etree.SubElement(inline, f"{{{WP}}}docPr")
    docPr.set("id", "900100")
    docPr.set("name", "FotoRosto")

    etree.SubElement(inline, f"{{{WP}}}cNvGraphicFramePr")

    graphic = etree.SubElement(inline, f"{{{A}}}graphic")
    graphicData = etree.SubElement(graphic, f"{{{A}}}graphicData")
    graphicData.set("uri", PIC)

    pic = etree.SubElement(graphicData, f"{{{PIC}}}pic")
    nvPicPr = etree.SubElement(pic, f"{{{PIC}}}nvPicPr")
    cNvPr = etree.SubElement(nvPicPr, f"{{{PIC}}}cNvPr")
    cNvPr.set("id", "900100")
    cNvPr.set("name", "FotoRosto")
    etree.SubElement(nvPicPr, f"{{{PIC}}}cNvPicPr")

    blipFill = etree.SubElement(pic, f"{{{PIC}}}blipFill")
    blip = etree.SubElement(blipFill, f"{{{A}}}blip")
    blip.set(f"{{{R}}}embed", rId)
    stretch = etree.SubElement(blipFill, f"{{{A}}}stretch")
    etree.SubElement(stretch, f"{{{A}}}fillRect")

    spPr = etree.SubElement(pic, f"{{{PIC}}}spPr")
    xfrm = etree.SubElement(spPr, f"{{{A}}}xfrm")
    off = etree.SubElement(xfrm, f"{{{A}}}off")
    off.set("x", "0")
    off.set("y", "0")
    ext = etree.SubElement(xfrm, f"{{{A}}}ext")
    ext.set("cx", "1080000")
    ext.set("cy", "1440000")
    prstGeom = etree.SubElement(spPr, f"{{{A}}}prstGeom")
    prstGeom.set("prst", "rect")
    etree.SubElement(prstGeom, f"{{{A}}}avLst")

    log.append("OK  [foto de rosto] inserida 3x4cm na celula reservada ao lado de DADOS DO(A) PERICIADO(A)")
    return True


def append_photo_annex(docx_path, photos, out_path, work_dir):
    """Usa python-docx para adicionar, ao final do documento, pagina(s) de
    anexo com as fotos fornecidas, dispostas lado a lado (2 por linha).
    Antes de inserir, cada foto e processada (recorte do fundo, correcao de
    perspectiva/rotacao e realce de nitidez/contraste, como um scanner) e
    encaixada preservando a proporcao original (sem esticar/distorcer).
    Mantem a mesma orientacao (retrato) do resto do laudo -- nao usa quebra
    de secao/paisagem, que em alguns leitores (Word) gerava uma pagina em
    branco antes do anexo."""
    import docx
    from docx.shared import Cm, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from PIL import Image

    d = docx.Document(docx_path)

    # remove paragrafos vazios sobrando no fim do documento
    body = d.element.body
    body_children = list(body)
    for child in reversed(body_children):
        if child.tag.endswith('}p'):
            has_text = "".join(t.text or "" for t in child.iter() if t.tag.endswith('}t')).strip()
            has_drawing = any(el.tag.endswith('}drawing') for el in child.iter())
            has_sectPr = any(el.tag.endswith('}sectPr') for el in child.iter())
            if not has_text and not has_drawing and not has_sectPr:
                body.remove(child)
                continue
        break

    processed = []
    scan_dir = Path(work_dir) / "anexo_processado"
    scan_dir.mkdir(parents=True, exist_ok=True)
    for i, photo in enumerate(photos):
        out_img = scan_dir / f"anexo_{i+1}.jpg"
        if scan_documento is not None:
            try:
                scan_documento.process(photo, out_img)
            except Exception:
                shutil.copy(photo, out_img)
        else:
            shutil.copy(photo, out_img)
        processed.append(out_img)

    d.add_page_break()
    h = d.add_paragraph()
    hr = h.add_run("ANEXO I - DOCUMENTOS APRESENTADOS NA PERÍCIA")
    hr.bold = True
    hr.font.size = Pt(14)

    # largura util da pagina (retrato) menos margens, dividida em 2 colunas
    section = d.sections[0]
    usable_w_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    col_w_cm = (usable_w_cm - 0.5) / 2  # 0.5cm de folga entre colunas
    MAX_W, MAX_H = col_w_cm, col_w_cm * 0.55  # mantem proporcao ~10x5,5

    def fit_size(img_path):
        with Image.open(img_path) as im:
            w, h = im.size
        aspect = w / h
        if aspect > (MAX_W / MAX_H):
            return Cm(MAX_W), Cm(MAX_W / aspect)
        else:
            return Cm(MAX_H * aspect), Cm(MAX_H)

    pairs = [processed[i:i + 2] for i in range(0, len(processed), 2)]
    for pair in pairs:
        table = d.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col in table.columns:
            col.width = Cm(usable_w_cm / 2)
        row = table.rows[0]
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            borders = docx.oxml.OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                el = docx.oxml.OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'nil')
                borders.append(el)
            tcPr.append(borders)
            p = cell.paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            if i < len(pair):
                width, height = fit_size(pair[i])
                run = p.add_run()
                run.add_picture(str(pair[i]), width=width, height=height)
        d.add_paragraph()  # espaco entre linhas de fotos

    d.save(out_path)


def process_case(modelo, dados_txt, fotos, output, foto_rosto=None, work_dir=None):
    """Processa um caso completo: preenche o modelo com os dados do txt,
    insere a foto de rosto (se houver) e anexa as fotos de documentos (se
    houver). Retorna a lista de linhas de log. Lanca excecao se algo
    impeditivo falhar (ex: modelo ou txt nao encontrado)."""
    work = Path(work_dir) if work_dir else Path("/home/claude/work/_tmp_preenchimento")
    if work.exists():
        shutil.rmtree(work)
    work.mkdir(parents=True)

    unpacked = work / "unpacked"
    unpacked.mkdir()
    with zipfile.ZipFile(modelo) as zf:
        zf.extractall(unpacked)

    data, name = parse_txt(dados_txt)

    log = []
    if name:
        log.append(f"Nome do(a) periciado(a) detectado: {name} (nao ha campo de nome no modelo; usar no nome do arquivo)")

    doc_xml_path = unpacked / "word" / "document.xml"
    tree = etree.parse(str(doc_xml_path))
    root = tree.getroot()

    fill_document(root, data, log)

    if foto_rosto:
        cropped_face = work / "foto_rosto_recortada.jpg"
        try:
            crop_to_ratio(foto_rosto, cropped_face, 3.0, 4.0)
            face_photo_path = cropped_face
        except Exception:
            face_photo_path = foto_rosto
        add_face_photo(root, unpacked, face_photo_path, log)

    tree.write(str(doc_xml_path), xml_declaration=True, encoding="UTF-8", standalone=True)

    intermediate = work / "intermediate.docx"
    rezip(unpacked, intermediate)

    final_path = Path(output)
    final_path.parent.mkdir(parents=True, exist_ok=True)
    if fotos:
        append_photo_annex(intermediate, fotos, final_path, work)
    else:
        shutil.copy(intermediate, final_path)

    return log


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("modelo")
    ap.add_argument("dados_txt")
    ap.add_argument("fotos", nargs="*")
    ap.add_argument("-o", "--output", required=True)
    ap.add_argument("--foto-rosto", dest="foto_rosto", default=None,
                     help="Foto de rosto do(a) periciado(a), inserida 3x4cm no topo da pagina 1")
    args = ap.parse_args()

    log = process_case(args.modelo, args.dados_txt, args.fotos, args.output, args.foto_rosto)

    print("\n=== LOG DE PREENCHIMENTO ===")
    for line in log:
        print(line)
    print(f"\nArquivo gerado: {args.output}")


if __name__ == "__main__":
    main()
