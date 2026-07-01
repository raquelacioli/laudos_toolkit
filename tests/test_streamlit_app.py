import tempfile
import unittest
from pathlib import Path

from docx import Document

from app import generate_docx_from_uploads


class StreamlitAppTests(unittest.TestCase):
    def test_generate_docx_from_uploads_creates_output(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            modelo = tmp / "modelo.docx"
            dados = tmp / "dados.txt"
            output = tmp / "saida.docx"

            doc = Document()
            doc.add_heading("Modelo de teste", level=1)
            doc.add_paragraph("DADOS DO(A) PERICIADO(A)")
            doc.add_paragraph("Idade:")
            doc.add_paragraph("3. ANAMNESE")
            doc.add_paragraph("4. EXAME FÍSICO E MENTAL")
            doc.save(modelo)

            dados.write_text(
                "FULANO DE TAL\nIDADE: 55 ANOS\nESCOLARIDADE: ENSINO MÉDIO COMPLETO\n"
                "ESTADO CIVIL: SOLTEIRO\nOCUPAÇÃO HABITUAL: PEDREIRO\n"
                "TEMPO DE AFASTAMENTO: HÁ 2 ANOS\n\nANAMNESE:\nTexto de teste.\n\n"
                "EXAME FÍSICO E MENTAL:\nTexto de teste.\n\nCONCLUSÃO:\nTexto de teste.\n",
                encoding="utf-8",
            )

            log = generate_docx_from_uploads(
                modelo_path=modelo,
                dados_txt_path=dados,
                fotos_paths=[],
                output_path=output,
                work_dir=tmp / "work",
            )

            self.assertTrue(output.exists())
            self.assertIsInstance(log, list)


if __name__ == "__main__":
    unittest.main()
